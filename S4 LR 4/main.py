from abc import ABC, abstractmethod
from docx import Document
from docx.shared import Pt


class ReportBuilder(ABC):
    @abstractmethod
    def add_header(self, title, theme, goal): pass

    @abstractmethod
    def add_paragraph(self, text): pass

    @abstractmethod
    def add_subheading(self, text, level): pass

    @abstractmethod
    def add_table(self, headers, rows): pass

    @abstractmethod
    def add_conclusion(self, text): pass

    @abstractmethod
    def save(self, filename): pass


class DocxReportBuilder(ReportBuilder):
    def __init__(self):
        self.document = Document()

    def add_header(self, title, theme, goal):
        self.document.add_heading(title, 0)
        
        p_theme = self.document.add_paragraph()
        p_theme.add_run(f"Тема работы: {theme}").font.size = Pt(16)
        
        p_goal = self.document.add_paragraph()
        p_goal.add_run(f"Цель работы: {goal}").font.size = Pt(16)

    def add_paragraph(self, text):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(14)

    def add_subheading(self, text, level):
        self.document.add_heading(text, level=level)

    def add_table(self, headers, rows):
        table = self.document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid' 
        
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            
        for i, row in enumerate(rows, start=1):
            for j, cell in enumerate(row):
                table.rows[i].cells[j].text = str(cell)

    def add_conclusion(self, text):
        self.document.add_heading("Вывод", level=1)
        self.add_paragraph(text)

    def save(self, filename):
        self.document.save(filename)
        self.document = Document()


class HTMLReportBuilder(ReportBuilder):
    def __init__(self):
        self.parts = []

    def add_header(self, title, theme, goal):
        self.parts.append(f"<h1>{title}</h1>")
        self.parts.append(f"<h2>Тема работы: {theme}</h2>")
        self.parts.append(f"<h2><b>Цель работы:</b> {goal}</h2>")

    def add_paragraph(self, text):
        self.parts.append(f"<p>{text}</p>")

    def add_subheading(self, text, level):
        tag = f"h{level+1}"
        self.parts.append(f"<{tag}>{text}</{tag}>")

    def add_table(self, headers, rows):
        html = "<table border='1' style='border-collapse: collapse; width: 100%;'>"
        html += "<tr style='background-color: #f2f2f2;'>"
        html += "".join(f"<th>{h}</th>" for h in headers) + "</tr>"
        for row in rows:
            html += "<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>"
        html += "</table>"
        self.parts.append(html)

    def add_conclusion(self, text):
        self.parts.append("<h2>Вывод</h2>")
        self.parts.append(f"<p>{text}</p>")

    def save(self, filename):
        full_html = f"<html><body style='font-family: Arial;'>\n" + "\n".join(self.parts) + "\n</body></html>"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(full_html)
        self.parts = []


class Director:
    def __init__(self, builder: ReportBuilder):
        self.builder = builder

    def build_report(self, data, filename):
        self.builder.add_header(data["title"], data["theme"], data["goal"])

        for block in data["body"]:
            if block["type"] == "text":
                self.builder.add_paragraph(block["content"])
            elif block["type"] == "heading_2":
                self.builder.add_subheading(block["content"], level=1)
            elif block["type"] == "heading_3":
                self.builder.add_subheading(block["content"], level=2)
            elif block["type"] == "table":
                self.builder.add_table(block["headers"], block["rows"])

        self.builder.add_conclusion(data["conclusion"])
        
        self.builder.save(filename)


report_data = {
    "title": "Лабораторная работа",
    "theme": "Доверительный интервал",
    "goal": "Изучить методы построения доверительных интервалов для среднего значения, доли и дисперсии генеральной совокупности",

    "body": [

        {"type": "heading_2", "content": "Задание 1."},
        {"type": "text", "content": "Постановка задачи: Азимут оси ВПП измерен в 4 приема. Результаты измерений представлены в таблице. Требуется оценить точность результата измерения."},

        {
            "type": "table",
            "headers": ["i", "xi", "xi - x", "(xi - x)^2"],
            "rows": [
                ["1", "25° 17' 30''", "-15''", "225"],
                ["2", "25° 17' 45''", "0''", "0"],
                ["3", "25° 18' 15''", "30''", "900"],
                ["4", "25° 17' 30''", "-15''", "225"]
            ]
        },

        {"type": "heading_3", "content": "Решение:"},
        {"type": "text", "content": "Начальные данные: n = 4, степень свободы = 3, доверительная вероятность = 0,95."},
        {"type": "text", "content": "Среднее значение x̄ = 25° 17' 45''."},
        {"type": "text", "content": "Дисперсия S2 = 450, стандартное отклонение S = 21,21."},
        {"type": "text", "content": "Среднеквадратическое отклонение центра распределения = 10,61."},
        {"type": "text", "content": "Коэффициент Стьюдента t = 3,18."},
        {"type": "text", "content": "Предельная погрешность ≈ 34."},
        {"type": "text", "content": "Доверительный интервал: от 25° 17' 11'' до 25° 18' 19''."},

        {"type": "heading_2", "content": "Задание 2."},
        {"type": "text", "content": "Постановка задачи: Проверен расход энергии в 10 квартирах. Данные: 125, 78, 102, 140, 90, 45, 50, 125, 115, 112."},
        {"type": "text", "content": "Определить доверительный интервал с надежностью 0,95."},

        {"type": "heading_3", "content": "Решение:"},
        {"type": "text", "content": "N = 70, n = 10, γ = 0,95."},
        {"type": "text", "content": "t = 2,26."},
        {"type": "text", "content": "Среднее значение x̄ = 98,2, стандартное отклонение S = 32,145."},
        {"type": "text", "content": "Предельная погрешность Δ = 21,27."},
        {"type": "text", "content": "Доверительный интервал: (76,93; 119,47)."},

        {"type": "heading_2", "content": "Задание 3."},
        {"type": "text", "content": "Постановка задачи: Определить объем выборки при N = 1000 и надежности 0,95."},

        {"type": "heading_3", "content": "Решение:"},
        {"type": "text", "content": "t = 1,96."},
        {"type": "text", "content": "Повторная выборка: n = 384."},
        {"type": "text", "content": "Бесповторная выборка: n = 277."},

        {"type": "heading_2", "content": "Задание 4."},
        {"type": "text", "content": "Постановка задачи: В партии 5000 изделий, проверено 400, из них 300 высшего сорта."},

        {"type": "heading_3", "content": "Решение:"},
        {"type": "text", "content": "t = 1,96, W = 0,75."},
        {"type": "text", "content": "Интервал (повторная): (0,7076; 0,7924)."},
        {"type": "text", "content": "Интервал (бесповторная): (0,7093; 0,7907)."},

        {"type": "heading_2", "content": "Задание 5."},
        {"type": "text", "content": "Постановка задачи: Найти объем выборки для 10000 банок при ошибке 0,05 и надежности 0,9995."},

        {"type": "heading_3", "content": "Решение:"},
        {"type": "text", "content": "t = 3,5, W = 0,5."},
        {"type": "text", "content": "Повторная выборка: n = 1225."},
        {"type": "text", "content": "Бесповторная выборка: n = 1091."},

        {"type": "heading_2", "content": "Задание 6."},
        {"type": "text", "content": "Постановка задачи: Найти доверительный интервал для среднего квадратического отклонения."},

        {"type": "heading_3", "content": "Решение:"},
        {"type": "text", "content": "Среднее значение = 0,3."},
        {"type": "text", "content": "Дисперсия S2 = 0,01158."},
        {"type": "text", "content": "Интервал для дисперсии: (0,006; 0,034)."},
        {"type": "text", "content": "Интервал для СКО: (0,077; 0,184)."},
    ],

    "conclusion": "В ходе лабораторной работы были построены доверительные интервалы для среднего значения, доли и среднего квадратического отклонения для повторной и бесповторной выборок."
}

if __name__ == "__main__":
    director = Director(DocxReportBuilder())
    director.build_report(report_data, "report.docx")

    director = Director(HTMLReportBuilder())
    director.build_report(report_data, "report.html")

    print("Отчёты успешно созданы!")